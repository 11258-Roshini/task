from flask import Flask, request, send_file, jsonify
from pathlib import Path
from io import BytesIO
import os
import logging
import html
import re

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib import colors
from PIL import Image as PILImage

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {'.docx', '.rtf', '.txt'}
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)



def allowed_file(filename):
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def safe_text(text):
    return html.escape(text or "").replace('\x00', '')


def get_image(image_stream):
    img = PILImage.open(image_stream)
    w, h = img.size
    aspect = h / w

    new_width = 5 * inch
    new_height = new_width * aspect

    image_stream.seek(0)
    return RLImage(image_stream, width=new_width, height=new_height)


def create_pdf_doc(buffer):
    return SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=0.5 * inch,
        leftMargin=0.5 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch
    )


#docx
def docx_to_story(input_path):
    doc = Document(input_path)
    story = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18)
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=11, leading=14)

    for para in doc.paragraphs:
        text = para.text.strip()

        # Extract images
        for run in para.runs:
            drawings = run._element.findall(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
            )
            for drawing in drawings:
                try:
                    blips = drawing.findall(
                        './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
                    )
                    for blip in blips:
                        embed_id = blip.get(
                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                        )
                        if embed_id and embed_id in doc.part.rels:
                            image_part = doc.part.rels[embed_id].target_part
                            image_stream = BytesIO(image_part.blob)

                            story.append(get_image(image_stream))
                            story.append(Spacer(1, 0.2 * inch))
                except Exception as e:
                    logger.error(f"Image error: {e}")

        if not text:
            story.append(Spacer(1, 0.1 * inch))
            continue

        text = safe_text(text)

        if 'Heading 1' in para.style.name:
            story.append(Paragraph(text, title_style))
        elif 'Heading' in para.style.name:
            story.append(Paragraph(text, heading_style))
        else:
            story.append(Paragraph(text, body_style))

        story.append(Spacer(1, 0.15 * inch))

    # Tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [safe_text(''.join(p.text for p in cell.paragraphs)) for cell in row.cells]
            table_data.append(row_data)

        if table_data:
            tbl = Table(table_data)
            tbl.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 0.2 * inch))

    return story


#rtf

def rtf_to_story(input_path):
    try:
        from striprtf.striprtf import rtf_to_text
        text = rtf_to_text(open(input_path).read())
    except:
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        text = re.sub(r'\\[a-z]+\d*\s?', '', content)
        text = re.sub(r'[{}]', '', text)

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=11)

    story = []
    for para in text.split('\n'):
        if para.strip():
            story.append(Paragraph(safe_text(para.strip()), body_style))
            story.append(Spacer(1, 0.1 * inch))

    return story


# txt

def txt_to_story(input_path):
    with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()

    styles = getSampleStyleSheet()
    text_style = ParagraphStyle('Text', parent=styles['Normal'], fontName='Courier')

    story = []
    for para in content.split('\n'):
        if para.strip():
            story.append(Paragraph(safe_text(para), text_style))
            story.append(Spacer(1, 0.1 * inch))

    return story


# route

@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400

    file = request.files['file']

    if not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, file.filename)
    output_path = os.path.join(OUTPUT_FOLDER, Path(file.filename).stem + '.pdf')

    try:
        file.save(input_path)

        pdf_doc = create_pdf_doc(output_path)

        ext = Path(file.filename).suffix.lower()

        if ext == '.docx':
            story = docx_to_story(input_path)
        elif ext == '.rtf':
            story = rtf_to_story(input_path)
        else:
            story = txt_to_story(input_path)

        pdf_doc.build(story)

        return send_file(output_path, as_attachment=True)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

    finally:
        if os.path.exists(input_path):
            os.remove(input_path)


# main

if __name__ == '__main__':
    app.run(debug=True)